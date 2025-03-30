import os
import shutil
from datetime import datetime
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Image as RLImage, Paragraph, Table, TableStyle, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import mm
from openpyxl import Workbook
from openpyxl.drawing.image import Image as XLImage
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from tkinter import messagebox
from PIL import Image, ImageTk
import openpyxl
import tkinter as tk
import re

def validar_nome(nome, max_length=100):
    if not nome or nome.strip() == "":
        return "O campo Nome é obrigatório!"
    if len(nome) > max_length:
        return f"O Nome deve ter no máximo {max_length} caracteres!"
    if not re.match(r"^[A-Za-z0-9À-ÿ\s]+$", nome):
        return "O Nome deve conter apenas letras, números e espaços!"
    return None

def validar_ano(ano):
    try:
        ano_int = int(ano)
        ano_atual = datetime.now().year
        if ano_int < 2000 or ano_int > ano_atual + 1:
            return f"O Ano deve estar entre 2000 e {ano_atual + 1}!"
        return None
    except ValueError:
        return "O Ano deve ser um número inteiro!"

def remover_foto(foto_path):
    try:
        if foto_path and os.path.exists(foto_path):
            os.remove(foto_path)
    except Exception as e:
        messagebox.showerror("Erro", f"Erro ao remover foto: {e}")

class SistemaController:
    def __init__(self, model, view):
        self.model = model
        self.view = view
        self.FONT = None
        self.BG_COLOR = None
        self.FG_COLOR = None
        self.backup_db()  # Criar backup ao iniciar

    def configurar_view(self, view):
        self.view = view
        self.FONT = view.FONT
        self.BG_COLOR = view.BG_COLOR
        self.FG_COLOR = view.FG_COLOR
        self.setup_menu()

    def setup_menu(self):
        menubar = self.view.root.nametowidget(self.view.root.cget("menu"))
        cadastro_menu = menubar.children["!menu"]
        cadastro_menu.entryconfig("Instituição", command=self.cadastro_instituicao)        
        cadastro_menu.entryconfig("Curso", command=self.cadastro_curso)
        cadastro_menu.entryconfig("Turma", command=self.cadastro_turma)
        cadastro_menu.entryconfig("Professor", command=self.cadastro_professor)
        cadastro_menu.entryconfig("Aluno", command=self.cadastro_aluno)

        consulta_menu = menubar.children["!menu2"]
        consulta_menu.entryconfig("Instituições", command=self.consulta_instituicoes)        
        consulta_menu.entryconfig("Cursos", command=self.consulta_cursos)
        consulta_menu.entryconfig("Turmas", command=self.consulta_turmas)
        consulta_menu.entryconfig("Professores", command=self.consulta_professores)
        consulta_menu.entryconfig("Alunos", command=self.consulta_alunos)

        carometro_menu = menubar.children["!menu3"]
        carometro_menu.add_command(label="Visualizar Carômetro", command=self.visualizar_carometro)
        carometro_menu.entryconfig("Exportar PDF", command=self.exportar_carometro_pdf)
        carometro_menu.entryconfig("Exportar Excel", command=self.exportar_carometro_excel)
        carometro_menu.entryconfig("Exportar Word", command=self.exportar_carometro_word)

    def fechar_conexao(self):
        """Fecha a conexão com o banco e cria um backup final."""
        try:
            self.backup_db()  # Criar backup ao encerrar
            self.model.close()
        except Exception as e:
            self.view.mostrar_erro(f"Erro ao fechar conexão com o banco: {e}")

    def backup_db(self):
        """Chama o método de backup do modelo."""
        try:
            self.model.backup_db()
        except Exception as e:
            self.view.mostrar_erro(f"Erro ao criar backup do banco: {e}")

    def validar_dados(self, tipo):
        """Valida os dados antes de salvar."""
        nome = self.entradas["nome"].get().strip()
        erro = validar_nome(nome)
        if erro:
            self.view.mostrar_aviso(erro)
            return False

        if tipo == "aluno":
            turma = self.entradas["turma"].get()
            if not turma or not turma.split(" - ")[0].isdigit():
                self.view.mostrar_aviso("O campo Turma é obrigatório e deve ser válido!")
                return False
            foto_path = self.entradas["foto"].get()
            if foto_path and not os.path.exists(foto_path):
                self.view.mostrar_aviso("O arquivo de foto selecionado não existe!")
                return False

        elif tipo == "professor":
            instituicao = self.entradas["instituição"].get()
            if not instituicao or not instituicao.split(" - ")[0].isdigit():
                self.view.mostrar_aviso("O campo Instituição é obrigatório e deve ser válido!")
                return False
            foto_path = self.entradas["foto"].get()
            if foto_path and not os.path.exists(foto_path):
                self.view.mostrar_aviso("O arquivo de foto selecionado não existe!")
                return False

        elif tipo == "curso":
            instituicao = self.entradas["instituição"].get()
            if not instituicao or not instituicao.split(" - ")[0].isdigit():
                self.view.mostrar_aviso("O campo Instituição é obrigatório e deve ser válido!")
                return False

        elif tipo == "turma":
            ano = self.entradas["ano"].get()
            curso = self.entradas["curso"].get()
            erro_ano = validar_ano(ano)
            if erro_ano:
                self.view.mostrar_aviso(erro_ano)
                return False
            if not curso or not curso.split(" - ")[0].isdigit():
                self.view.mostrar_aviso("O campo Curso é obrigatório e deve ser válido!")
                return False

        return True

    # Consultas
    def consulta_instituicoes(self):
        def atualizar(): self.atualizar_tabela("instituicoes", ["id", "nome"])
        def duplo_clique(event): self.on_double_click(event, self.cadastro_instituicao)
        self.tree, self.filtros = self.view.consulta_generica("Consulta de Instituições", ["id", "nome"], atualizar, duplo_clique)
        atualizar()

    def consulta_professores(self):
        def atualizar(): self.atualizar_tabela("professores", ["id", "nome", "instituicao"])
        def duplo_clique(event): self.on_double_click(event, self.cadastro_professor)
        self.tree, self.filtros = self.view.consulta_generica("Consulta de Professores", ["id", "nome", "instituicao"], atualizar, duplo_clique)
        atualizar()

    def consulta_cursos(self):
        def atualizar(): self.atualizar_tabela("cursos", ["id", "nome", "instituicao"])
        def duplo_clique(event): self.on_double_click(event, self.cadastro_curso)
        self.tree, self.filtros = self.view.consulta_generica("Consulta de Cursos", ["id", "nome", "instituicao"], atualizar, duplo_clique)
        atualizar()

    def consulta_turmas(self):
        def atualizar(): self.atualizar_tabela("turmas", ["id", "nome", "ano", "curso"])
        def duplo_clique(event): self.on_double_click(event, self.cadastro_turma)
        self.tree, self.filtros = self.view.consulta_generica("Consulta de Turmas", ["id", "nome", "ano", "curso"], atualizar, duplo_clique)
        atualizar()

    def consulta_alunos(self):
        def atualizar(): self.atualizar_tabela("alunos", ["id", "nome", "turma", "curso", "instituicao"])
        def duplo_clique(event): self.on_double_click(event, self.cadastro_aluno)
        self.tree, self.filtros = self.view.consulta_generica("Consulta de Alunos", ["id", "nome", "turma", "curso", "instituicao"], atualizar, duplo_clique)
        atualizar()

    def atualizar_tabela(self, tipo, colunas):
        try:
            filtros = {col: self.filtros[col].get() for col in colunas[1:]}
            metodo = getattr(self.model, f"consulta_{tipo}")
            dados = metodo(filtros)
            self.view.atualizar_tabela(self.tree, dados)
        except Exception as e:
            self.view.mostrar_erro(f"Erro ao consultar {tipo}: {e}")

    def on_double_click(self, event, callback):
        item = self.tree.selection()
        if item:
            id_value = self.tree.item(item[0], "values")[0]
            print(f"Debug: ID selecionado = {id_value}")
            callback(int(id_value))
        else:
            print("Debug: Nenhum item selecionado na tabela")

    # Cadastros
    def cadastro_instituicao(self, id=None):
        campos = [("Nome", "entry", None)]
        def salvar(): self.salvar_instituicao(id)
        def excluir(): self.excluir_instituicao(id)
        self.entradas = self.view.cadastro_generico(f"Cadastro de Instituição{' - Editar' if id else ''}", campos, salvar, excluir if id else None)
        if id:
            dados = self.model.executar_query("SELECT nome FROM instituicao WHERE id = ?", (id,), fetch=True)
            if dados:
                self.entradas["nome"].insert(0, dados[0][0])

    def cadastro_professor(self, id=None):
        instituicoes = [f"{i[0]} - {i[1]}" for i in self.model.carregar_instituicoes()]
        campos = [("Nome", "entry", None), ("Instituição", "combo", instituicoes), ("Foto", "foto", None)]
        def salvar(): self.salvar_professor(id)
        def excluir(): self.excluir_professor(id)
        self.entradas = self.view.cadastro_generico(f"Cadastro de Professor{' - Editar' if id else ''}", campos, salvar, excluir if id else None)
        if id:
            dados = self.model.executar_query("SELECT nome, instituicao_id, foto FROM professor WHERE id = ?", (id,), fetch=True)
            print(f"Debug: Dados retornados para ID {id} = {dados}")
            if dados:
                self.entradas["nome"].delete(0, tk.END)
                self.entradas["nome"].insert(0, dados[0][0])
                instituicao_nome = self.model.executar_query("SELECT nome FROM instituicao WHERE id = ?", (dados[0][1],), fetch=True)
                print(f"Debug: Instituição retornada = {instituicao_nome}")
                if instituicao_nome:
                    self.entradas["instituição"].set(f"{dados[0][1]} - {instituicao_nome[0][0]}")
                else:
                    self.view.mostrar_erro(f"Instituição associada ao professor (ID {dados[0][1]}) não encontrada!")
                    self.entradas["instituição"].set("")
                base_dir = os.path.dirname(__file__)
                foto_path = os.path.join(base_dir, dados[0][2]) if dados[0][2] else None
                print(f"Debug: Caminho da foto = {foto_path}")
                if foto_path and os.path.exists(foto_path):
                    self.entradas["foto"].delete(0, tk.END)
                    self.entradas["foto"].insert(0, foto_path)
                    try:
                        img = Image.open(foto_path)
                        img = img.resize((150, 150), Image.Resampling.LANCZOS)
                        foto = ImageTk.PhotoImage(img)
                        self.entradas["foto_label"].config(image=foto)
                        self.entradas["foto_label"].image = foto
                    except Exception as e:
                        self.view.mostrar_erro(f"Não foi possível carregar a imagem: {e}")
                else:
                    self.entradas["foto"].delete(0, tk.END)
                    self.entradas["foto_label"].config(image="")
            else:
                self.view.mostrar_erro(f"Professor com ID {id} não encontrado!")
                self.view.tela_inicial()

    def cadastro_curso(self, id=None, tree=None):
        instituicoes = [f"{i[0]} - {i[1]}" for i in self.model.carregar_instituicoes()]
        campos = [("Nome", "entry", None), ("Instituição", "combo", instituicoes)]
        def salvar(): self.salvar_curso(id)
        def excluir(): self.excluir_curso(id)
        self.entradas = self.view.cadastro_generico(f"Cadastro de Curso{' - Editar' if id else ''}", campos, salvar, excluir if id else None)
        self.tree = tree
        if id:
            dados = self.model.executar_query("SELECT nome, instituicao_id FROM curso WHERE id = ?", (id,), fetch=True)
            if dados:
                self.entradas["nome"].insert(0, dados[0][0])
                self.entradas["instituição"].set(f"{dados[0][1]} - {self.model.executar_query('SELECT nome FROM instituicao WHERE id = ?', (dados[0][1],), fetch=True)[0][0]}")

    def cadastro_turma(self, id=None, tree=None):
        cursos = [f"{c[0]} - {c[1]}" for c in self.model.carregar_cursos()]
        campos = [
            ("Nome", "entry", None),
            ("Ano", "entry", None),
            ("Curso", "combo", cursos)
        ]
        def salvar(): self.salvar_turma(id)
        def excluir(): self.excluir_turma(id)
        self.entradas = self.view.cadastro_generico(f"Cadastro de Turma{' - Editar' if id else ''}", campos, salvar, excluir if id else None)
        self.tree = tree
        if id:
            dados = self.model.executar_query("SELECT nome, ano, curso_id FROM turma WHERE id = ?", (id,), fetch=True)
            if dados:
                self.entradas["nome"].insert(0, dados[0][0])
                self.entradas["ano"].insert(0, dados[0][1])
                self.entradas["curso"].set(f"{dados[0][2]} - {self.model.executar_query('SELECT nome FROM curso WHERE id = ?', (dados[0][2],), fetch=True)[0][0]}")
        else:
            self.entradas["ano"].insert(0, str(datetime.now().year))

    def cadastro_aluno(self, id=None):
        turmas = [f"{t[0]} - {t[1]}" for t in self.model.carregar_turmas()]
        campos = [("Nome", "entry", None), ("Turma", "combo", turmas), ("Foto", "foto", None)]
        def salvar(): self.salvar_aluno(id)
        def excluir(): self.excluir_aluno(id)
        self.entradas = self.view.cadastro_generico(f"Cadastro de Aluno{' - Editar' if id else ''}", campos, salvar, excluir if id else None)
        if id:
            dados = self.model.executar_query("SELECT nome, turma_id, foto FROM aluno WHERE id = ?", (id,), fetch=True)
            print(f"Debug: Dados retornados para ID {id} = {dados}")
            if dados:
                self.entradas["nome"].delete(0, tk.END)
                self.entradas["nome"].insert(0, dados[0][0])
                turma_nome = self.model.executar_query("SELECT nome FROM turma WHERE id = ?", (dados[0][1],), fetch=True)
                print(f"Debug: Turma retornada = {turma_nome}")
                if turma_nome:
                    self.entradas["turma"].set(f"{dados[0][1]} - {turma_nome[0][0]}")
                else:
                    self.view.mostrar_erro(f"Turma associada ao aluno (ID {dados[0][1]}) não encontrada!")
                    self.entradas["turma"].set("")
                base_dir = os.path.dirname(__file__)
                foto_path = os.path.join(base_dir, dados[0][2]) if dados[0][2] else None
                print(f"Debug: Caminho da foto = {foto_path}")
                if foto_path and os.path.exists(foto_path):
                    self.entradas["foto"].delete(0, tk.END)
                    self.entradas["foto"].insert(0, foto_path)
                    try:
                        img = Image.open(foto_path)
                        img = img.resize((150, 150), Image.Resampling.LANCZOS)
                        foto = ImageTk.PhotoImage(img)
                        self.entradas["foto_label"].config(image=foto)
                        self.entradas["foto_label"].image = foto
                    except Exception as e:
                        self.view.mostrar_erro(f"Não foi possível carregar a imagem: {e}")
                else:
                    self.entradas["foto"].delete(0, tk.END)
                    self.entradas["foto_label"].config(image="")
            else:
                self.view.mostrar_erro(f"Aluno com ID {id} não encontrado!")
                self.view.tela_inicial()

    def salvar_instituicao(self, id):        
        try:
            if not self.validar_dados("instituicao"):
                return
            nome = self.entradas["nome"].get()
            self.model.salvar_instituicao(id, nome)
            self.view.mostrar_mensagem("Sucesso", f"Instituição {'atualizada' if id else 'salva'} com sucesso!")
            self.view.tela_inicial()
        except Exception as e:
            self.view.mostrar_erro(f"Erro ao salvar instituição: {e}")

    def salvar_professor(self, id):
        try:
            if not self.validar_dados("professor"):
                return
            nome = self.entradas["nome"].get()
            instituicao = self.entradas["instituição"].get().split(" - ")[0]
            foto_path = self.entradas["foto"].get()

            # Carregar o registro atual do banco, se for uma edição
            foto_antiga = None
            nome_antigo = None
            if id:
                dados = self.model.executar_query("SELECT nome, instituicao_id, foto FROM professor WHERE id = ?", (id,), fetch=True)
                if dados:
                    nome_antigo = dados[0][0]
                    foto_antiga = dados[0][2] if dados[0][2] else None
                    print(f"Debug: Foto antiga carregada do banco = {foto_antiga}, Nome antigo = {nome_antigo}")

            # Determinar se uma nova foto foi selecionada
            foto_path_final = None
            if foto_path and foto_path.strip():  # Verificar se foto_path não é uma string vazia
                if foto_antiga:
                    foto_antiga_absoluta = os.path.join(os.path.dirname(__file__), foto_antiga)
                    # Comparar o caminho absoluto da foto_path com o caminho da foto_antiga
                    if os.path.abspath(foto_path) == os.path.abspath(foto_antiga_absoluta):
                        print("Debug: O caminho da foto não mudou, nenhuma nova foto selecionada")
                        foto_path_final = None  # Não há nova foto, apenas renomear a antiga se necessário
                    else:
                        print("Debug: Nova foto selecionada, usando foto_path")
                        foto_path_final = foto_path
                else:
                    print("Debug: Nova foto selecionada, usando foto_path")
                    foto_path_final = foto_path
            else:
                print("Debug: Nenhuma nova foto selecionada (foto_path vazio)")
                foto_path_final = None

            # Salvar o registro no banco primeiro, sem a foto
            new_id = self.model.salvar_professor(id, nome, int(instituicao), None)
            print(f"Debug: Novo ID gerado/atualizado = {new_id}")
            self.model.commit()

            # Mover a foto (ou renomear, se necessário) usando o new_id
            foto = self.mover_foto(foto_path_final, "professores", new_id, nome, nome_antigo, foto_antiga) if (foto_path_final or foto_antiga) else None
            print(f"Debug: Caminho da foto após mover_foto = {foto}")

            # Atualizar o registro com o caminho da foto
            if foto:
                self.model.salvar_professor(new_id, nome, int(instituicao), foto)
                self.model.commit()

            self.view.mostrar_mensagem("Sucesso", f"Professor {'atualizado' if id else 'salvo'} com sucesso!")
            self.view.tela_inicial()
        except Exception as e:
            self.view.mostrar_erro(f"Erro ao salvar professor: {e}")

    def salvar_curso(self, id):
        try:
            if not self.validar_dados("curso"):
                return
            nome = self.entradas["nome"].get()
            instituicao = self.entradas["instituição"].get().split(" - ")[0]
            self.model.salvar_curso(id, nome, int(instituicao))
            self.view.mostrar_mensagem("Sucesso", f"Curso {'atualizado' if id else 'salvo'} com sucesso!")
            self.view.tela_inicial()
        except Exception as e:
            self.view.mostrar_erro(f"Erro ao salvar curso: {e}")

    def salvar_turma(self, id):
        try:
            if not self.validar_dados("turma"):
                return
            nome = self.entradas["nome"].get()
            ano = self.entradas["ano"].get()
            curso = self.entradas["curso"].get().split(" - ")[0]
            self.model.salvar_turma(id, nome, int(ano), int(curso))
            self.view.mostrar_mensagem("Sucesso", f"Turma {'atualizada' if id else 'salva'} com sucesso!")
            self.view.tela_inicial()
        except Exception as e:
            self.view.mostrar_erro(f"Erro ao salvar turma: {e}")

    def salvar_aluno(self, id):
        try:
            if not self.validar_dados("aluno"):
                return
            nome = self.entradas["nome"].get()
            turma = self.entradas["turma"].get().split(" - ")[0]
            foto_path = self.entradas["foto"].get()

            # Carregar o registro atual do banco, se for uma edição
            foto_antiga = None
            nome_antigo = None
            if id:
                dados = self.model.executar_query("SELECT nome, turma_id, foto FROM aluno WHERE id = ?", (id,), fetch=True)
                if dados:
                    nome_antigo = dados[0][0]
                    foto_antiga = dados[0][2] if dados[0][2] else None
                    print(f"Debug: Foto antiga carregada do banco = {foto_antiga}, Nome antigo = {nome_antigo}")

            # Determinar se uma nova foto foi selecionada
            foto_path_final = None
            if foto_path and foto_path.strip():  # Verificar se foto_path não é uma string vazia
                if foto_antiga:
                    foto_antiga_absoluta = os.path.join(os.path.dirname(__file__), foto_antiga)
                    # Comparar o caminho absoluto da foto_path com o caminho da foto_antiga
                    if os.path.abspath(foto_path) == os.path.abspath(foto_antiga_absoluta):
                        print("Debug: O caminho da foto não mudou, nenhuma nova foto selecionada")
                        foto_path_final = None  # Não há nova foto, apenas renomear a antiga se necessário
                    else:
                        print("Debug: Nova foto selecionada, usando foto_path")
                        foto_path_final = foto_path
                else:
                    print("Debug: Nova foto selecionada, usando foto_path")
                    foto_path_final = foto_path
            else:
                print("Debug: Nenhuma nova foto selecionada (foto_path vazio)")
                foto_path_final = None

            # Salvar o registro no banco primeiro, sem a foto
            new_id = self.model.salvar_aluno(id, nome, int(turma), None)
            print(f"Debug: Novo ID gerado/atualizado = {new_id}")
            self.model.commit()

            # Mover a foto (ou renomear, se necessário) usando o new_id
            foto = self.mover_foto(foto_path_final, "alunos", new_id, nome, nome_antigo, foto_antiga) if (foto_path_final or foto_antiga) else None
            print(f"Debug: Caminho da foto após mover_foto = {foto}")

            # Atualizar o registro com o caminho da foto
            if foto:
                self.model.salvar_aluno(new_id, nome, int(turma), foto)
                self.model.commit()

            self.view.mostrar_mensagem("Sucesso", f"Aluno {'atualizado' if id else 'salvo'} com sucesso!")
            self.view.tela_inicial()
        except Exception as e:
            self.view.mostrar_erro(f"Erro ao salvar aluno: {e}")

    def excluir_instituicao(self, id):
        try:
            if not id:
                self.view.mostrar_aviso("Nenhuma instituição selecionada para excluir!")
                return
            if not messagebox.askyesno("Confirmação", "Tem certeza que deseja excluir esta instituição?"):
                return
            self.model.excluir_registro("instituicao", id)
            self.view.mostrar_mensagem("Sucesso", "Instituição excluída com sucesso!")
            self.consulta_instituicoes()
        except Exception as e:
            self.view.mostrar_erro(f"Erro ao excluir instituição: {e}")

    def excluir_professor(self, id):
        try:
            if not id:
                self.view.mostrar_aviso("Nenhum professor selecionado para excluir!")
                return
            if not messagebox.askyesno("Confirmação", "Tem certeza que deseja excluir este professor?"):
                return
            dados = self.model.executar_query("SELECT foto FROM professor WHERE id = ?", (id,), fetch=True)
            if dados and dados[0][0]:
                foto_path = os.path.join(os.path.dirname(__file__), dados[0][0])
                remover_foto(foto_path)
            self.model.excluir_registro("professor", id)
            self.view.mostrar_mensagem("Sucesso", "Professor excluído com sucesso!")
            self.consulta_professores()
        except Exception as e:
            self.view.mostrar_erro(f"Erro ao excluir professor: {e}")

    def excluir_curso(self, id):
        try:
            if not id:
                self.view.mostrar_aviso("Nenhum curso selecionado para excluir!")
                return
            if not messagebox.askyesno("Confirmação", "Tem certeza que deseja excluir este curso?"):
                return
            self.model.excluir_registro("curso", id)
            self.view.mostrar_mensagem("Sucesso", "Curso excluído com sucesso!")
            self.consulta_cursos()
        except Exception as e:
            self.view.mostrar_erro(f"Erro ao excluir curso: {e}")

    def excluir_turma(self, id):
        try:
            if not id:
                self.view.mostrar_aviso("Nenhuma turma selecionada para excluir!")
                return
            if not messagebox.askyesno("Confirmação", "Tem certeza que deseja excluir esta turma?"):
                return
            self.model.excluir_registro("turma", id)
            self.view.mostrar_mensagem("Sucesso", "Turma excluída com sucesso!")
            self.consulta_turmas()
        except Exception as e:
            self.view.mostrar_erro(f"Erro ao excluir turma: {e}")

    def excluir_aluno(self, id):
        try:
            if not id:
                self.view.mostrar_aviso("Nenhum aluno selecionado para excluir!")
                return
            if not messagebox.askyesno("Confirmação", "Tem certeza que deseja excluir este aluno?"):
                return
            dados = self.model.executar_query("SELECT foto FROM aluno WHERE id = ?", (id,), fetch=True)
            if dados and dados[0][0]:
                foto_path = os.path.join(os.path.dirname(__file__), dados[0][0])
                remover_foto(foto_path)
            self.model.excluir_registro("aluno", id)
            self.view.mostrar_mensagem("Sucesso", "Aluno excluído com sucesso!")
            self.consulta_alunos()
        except Exception as e:
            self.view.mostrar_erro(f"Erro ao excluir aluno: {e}")

    def mover_foto(self, foto_path, tipo, id, nome, nome_antigo=None, foto_antiga=None):
        if id is None:
            raise ValueError("ID não pode ser None ao mover a foto!")
        print(f"Debug: Movendo foto - ID={id}, Nome={nome}, Nome Antigo={nome_antigo}, Tipo={tipo}, Foto Path={foto_path}, Foto Antiga={foto_antiga}")
        imagens_dir = os.path.join("imagens", tipo)
        os.makedirs(os.path.join(os.path.dirname(__file__), imagens_dir), exist_ok=True)
        nome_arquivo = f"{str(id).zfill(5)}-{nome.replace(' ', '_')}.jpg"
        print(f"Debug: Nome do arquivo gerado = {nome_arquivo}")
        caminho_relativo = os.path.join(imagens_dir, nome_arquivo)
        caminho_absoluto = os.path.join(os.path.dirname(__file__), caminho_relativo)

        # Se o nome foi alterado e existe uma foto antiga, renomear a foto
        if nome_antigo and nome != nome_antigo and foto_antiga:
            nome_arquivo_antigo = f"{str(id).zfill(5)}-{nome_antigo.replace(' ', '_')}.jpg"
            caminho_antigo_relativo = os.path.join(imagens_dir, nome_arquivo_antigo)
            caminho_antigo_absoluto = os.path.join(os.path.dirname(__file__), caminho_antigo_relativo)
            if os.path.exists(caminho_antigo_absoluto):
                try:
                    # Se uma nova foto foi selecionada, remover a antiga
                    if foto_path and os.path.exists(foto_path):
                        os.remove(caminho_antigo_absoluto)
                        print(f"Debug: Foto antiga removida devido a nova foto: {caminho_antigo_absoluto}")
                    else:
                        # Se não há nova foto, apenas renomear a antiga
                        os.rename(caminho_antigo_absoluto, caminho_absoluto)
                        print(f"Debug: Foto renomeada de {caminho_antigo_absoluto} para {caminho_absoluto}")
                        return caminho_relativo
                except Exception as e:
                    self.view.mostrar_erro(f"Erro ao renomear/remover foto antiga: {e}")
                    return None
            else:
                print(f"Debug: Foto antiga não encontrada: {caminho_antigo_absoluto}")

        # Se uma nova foto foi selecionada, copiar ou renomear para o novo local
        if foto_path and os.path.exists(foto_path):
            print(f"Debug: Novo caminho absoluto da foto = {os.path.abspath(foto_path)}")
            # Remover qualquer foto existente no caminho_absoluto (exceto se for a mesma que foto_path)
            if os.path.exists(caminho_absoluto):
                try:
                    os.remove(caminho_absoluto)
                    print(f"Debug: Foto existente removida: {caminho_absoluto}")
                except Exception as e:
                    self.view.mostrar_erro(f"Erro ao remover foto existente: {e}")
                    return None

            # Copiar ou renomear a nova foto
            try:
                if os.path.abspath(foto_path).startswith(os.path.abspath(imagens_dir)):
                    os.rename(foto_path, caminho_absoluto)
                    print(f"Debug: Foto renomeada de {foto_path} para {caminho_absoluto}")
                else:
                    shutil.copy(foto_path, caminho_absoluto)
                    print(f"Debug: Nova foto copiada para {caminho_absoluto}")
            except Exception as e:
                self.view.mostrar_erro(f"Erro ao mover/renomear foto: {e}")
                return None

            return caminho_relativo

        # Se não há nova foto, mas existe uma foto renomeada ou antiga, retornar o caminho relativo
        if os.path.exists(caminho_absoluto):
            print(f"Debug: Foto já existe no caminho correto, retornando {caminho_relativo}")
            return caminho_relativo

        print(f"Debug: Nenhuma foto para processar, retornando None")
        return None

    # Exportações
    def exportar_carometro_pdf(self):
        turmas = ["Todas"] + [f"{t[0]} - {t[1]}" for t in self.model.carregar_turmas()]
        combo = self.view.exportar_carometro("PDF", self.exportar_pdf)
        combo["values"] = turmas
        combo.set("Todas")

    def exportar_carometro_excel(self):
        turmas = ["Todas"] + [f"{t[0]} - {t[1]}" for t in self.model.carregar_turmas()]
        combo = self.view.exportar_carometro("Excel", self.exportar_excel)
        combo["values"] = turmas
        combo.set("Todas")

    def exportar_carometro_word(self):
        turmas = ["Todas"] + [f"{t[0]} - {t[1]}" for t in self.model.carregar_turmas()]
        combo = self.view.exportar_carometro("Word", self.exportar_word)
        combo["values"] = turmas
        combo.set("Todas")

    def visualizar_carometro(self):
        turma_combo, scrollable_frame = self.view.visualizar_carometro(self.atualizar_carometro)
        self.turma_combo = turma_combo
        self.scrollable_frame = scrollable_frame

    def atualizar_carometro(self):
        try:
            print("Debug: Iniciando tarefa de atualização do carômetro")
            turma_str = self.turma_combo.get()
            print(f"Debug: turma_str = {turma_str}")
            if not turma_str:
                print("Debug: Nenhuma turma selecionada")
                self.view.fechar_carregando()
                return
            turma_id = None if turma_str == "Todas" else int(turma_str.split(" - ")[0])
            print(f"Debug: turma_id = {turma_id}")
            alunos = self.model.carregar_alunos_por_turma(turma_id)
            print(f"Debug: Alunos retornados = {alunos}")
            
            base_dir = os.path.dirname(__file__)
            dados = []
            for aluno in alunos:
                print(f"Debug: Processando aluno = {aluno}")                    
                foto_path = os.path.join(base_dir, aluno[5]) if aluno[5] else os.path.join(base_dir, "imagens", "00000-sem imagem.jpg")
                if os.path.exists(foto_path):
                    img = Image.open(foto_path)
                    img = img.resize((100, 100), Image.Resampling.LANCZOS)
                    photo = ImageTk.PhotoImage(img)
                else:
                    photo = None
                dados.append((aluno[1], photo, foto_path))
                
            print(f"Debug: Dados preparados = {len(dados)} itens")
            self.atualizar_interface_carometro(dados)
        except Exception as e:
            erro = e
            print(f"Debug: Erro capturado em tarefa: {erro}")
            self.view.fechar_carregando()
            self.view.mostrar_erro(f"Erro ao atualizar carômetro: {erro}")

    def atualizar_interface_carometro(self, dados):
        try:
            print("Debug: Iniciando atualização da interface")
            for widget in self.scrollable_frame.winfo_children():
                widget.destroy()
            
            for i, (nome, photo, foto_path) in enumerate(dados):
                aluno_frame = tk.Frame(self.scrollable_frame, bg=self.BG_COLOR, bd=2, relief="groove")
                aluno_frame.grid(row=i//6, column=i%6, padx=5, pady=5, sticky="nsew")
                
                if photo:
                    tk.Label(aluno_frame, image=photo, bg=self.BG_COLOR).pack(pady=5)
                    aluno_frame.image = photo
                else:
                    tk.Label(aluno_frame, text="Sem Foto", font=self.FONT, bg=self.BG_COLOR, fg=self.FG_COLOR).pack(pady=5)
                tk.Label(aluno_frame, text=nome, font=self.FONT, bg=self.BG_COLOR, fg=self.FG_COLOR).pack(pady=5)
            
            print("Debug: Interface atualizada com sucesso")
            self.view.fechar_carregando()
        except Exception as e:
            print(f"Debug: Erro ao atualizar interface: {e}")
            self.view.fechar_carregando()
            self.view.mostrar_erro(f"Erro ao atualizar interface do carômetro: {e}")

    def exportar_pdf(self):
        try:
            print("Debug: Iniciando exportar_pdf")
            turma_str = self.view.current_frame.winfo_children()[2].get()
            print(f"Debug: turma_str = {turma_str}")
            if not turma_str:
                self.view.fechar_carregando()
                self.view.mostrar_aviso("Selecione uma turma!")
                return
            turma_id = None if turma_str == "Todas" else int(turma_str.split(" - ")[0])
            print(f"Debug: turma_id = {turma_id}")
            alunos = self.model.carregar_alunos_por_turma(turma_id)
            print(f"Debug: Alunos retornados = {alunos}")
            if not alunos:
                self.view.fechar_carregando()
                self.view.mostrar_mensagem("Info", "Nenhum aluno encontrado para a turma selecionada.")
                return

            agora = datetime.now()
            filtro = "todas" if turma_str == "Todas" else turma_str.split(" - ")[1].replace(" ", "_").lower()
            nome_arquivo = f"carometro-{filtro}-{agora.strftime('%Y-%m-%d_%H-%M-%S')}.pdf"
            documentos_dir = os.path.join(os.path.dirname(__file__), "documentos")
            os.makedirs(documentos_dir, exist_ok=True)
            pdf_path = os.path.join(documentos_dir, nome_arquivo)

            pdf = SimpleDocTemplate(pdf_path, pagesize=A4, leftMargin=9*mm, rightMargin=9*mm, topMargin=20*mm, bottomMargin=30*mm)
            elements = []
            styles = getSampleStyleSheet()
            style = styles["Normal"]
            style.fontSize = 8
            style.alignment = 1

            def build_pdf_footer(canvas, doc):
                canvas.saveState()
                page_number = canvas.getPageNumber()
                data_hora = agora.strftime("%d/%m/%Y às %H:%M")
                footer_text = f"Página {page_number}   |   Documento gerado em {data_hora}"
                canvas.setFont("Helvetica", 8)
                canvas.drawCentredString(A4[0] / 2, 15, footer_text)
                canvas.restoreState()

            elements.append(Paragraph(f"Carômetro - {turma_str}", styles["Heading1"]))
            elements.append(Spacer(1, 12))

            elementos_linha = []
            base_dir = os.path.dirname(__file__)
            for i, aluno in enumerate(alunos):
                print(f"Debug: Processando aluno = {aluno}")                    
                foto_path = os.path.join(base_dir, aluno[5]) if aluno[5] else os.path.join(base_dir, "imagens", "00000-sem imagem.jpg")
                print(f"Debug: foto_path = {foto_path}")
                primeiro_nome = aluno[1].split(" ")[0]
                if os.path.exists(foto_path):
                    try:
                        img = RLImage(foto_path, width=28*mm, height=28*mm)
                    except Exception as e:
                        print(f"Debug: Erro ao carregar foto {foto_path}: {e}")
                        img = Paragraph("Sem Foto", style)
                else:
                    print(f"Debug: Foto não encontrada: {foto_path}")
                    img = Paragraph("Sem Foto", style)
                nome = Paragraph(primeiro_nome, style)
                elementos_linha.append([img, nome])

                if (i + 1) % 6 == 0 or i == len(alunos) - 1:
                    num_elementos = len(elementos_linha)
                    num_linhas = (num_elementos + 5) // 6
                    if num_elementos > 0:
                        table_data = []
                        for r in range(num_linhas):
                            inicio = r * 6
                            fim = min(inicio + 6, num_elementos)
                            linha = elementos_linha[inicio:fim]
                            while len(linha) < 6:
                                linha.append([Paragraph("", style), Paragraph("", style)])
                            table_data.append(linha)

                        table = Table(table_data, colWidths=[32*mm]*6, rowHeights=[32*mm]*num_linhas)
                        table.setStyle(TableStyle([
                            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                            ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
                            ('FONTSIZE', (0, 0), (-1, -1), 8),
                        ]))
                        elements.append(table)
                        elementos_linha = []
                        if (i + 1) % 42 == 0 and i != len(alunos) - 1:
                            elements.append(PageBreak())

            pdf.build(elements, onFirstPage=build_pdf_footer, onLaterPages=build_pdf_footer)
            self.view.fechar_carregando()
            self.view.mostrar_mensagem("Sucesso", f"Carômetro exportado como {nome_arquivo} na pasta 'documentos'!")
            os.startfile(pdf_path)
            self.view.tela_inicial()
        except Exception as e:
            self.view.fechar_carregando()
            self.view.mostrar_erro(f"Erro ao exportar PDF: {e}")

    def exportar_excel(self):
        try:
            turma_str = self.view.current_frame.winfo_children()[2].get()
            if not turma_str:
                self.view.fechar_carregando()
                self.view.mostrar_aviso("Selecione uma turma!")
                return
            turma_id = None if turma_str == "Todas" else int(turma_str.split(" - ")[0])
            alunos = self.model.carregar_alunos_por_turma(turma_id)
            if not alunos:
                self.view.fechar_carregando()
                self.view.mostrar_mensagem("Info", "Nenhum aluno encontrado para a turma selecionada.")
                return

            agora = datetime.now()
            filtro = "todas" if turma_str == "Todas" else turma_str.split(" - ")[1].replace(" ", "_").lower()
            nome_arquivo = f"carometro-{filtro}-{agora.strftime('%Y-%m-%d_%H-%M-%S')}.xlsx"
            documentos_dir = os.path.join(os.path.dirname(__file__), "documentos")
            os.makedirs(documentos_dir, exist_ok=True)
            excel_path = os.path.join(documentos_dir, nome_arquivo)

            wb = Workbook()
            ws = wb.active
            ws.title = "Carômetro"
            ws.append(["ID", "Nome", "Turma", "Curso", "Instituição", "Foto"])

            ws.column_dimensions['A'].width = 10
            ws.column_dimensions['B'].width = 20
            ws.column_dimensions['C'].width = 15
            ws.column_dimensions['D'].width = 15
            ws.column_dimensions['E'].width = 20
            ws.column_dimensions['F'].width = 15

            temp_dir = os.path.join(os.path.dirname(__file__), "temp")
            os.makedirs(temp_dir, exist_ok=True)

            base_dir = os.path.dirname(__file__)
            for i, aluno in enumerate(alunos, start=2):
                ws.cell(row=i, column=1, value=aluno[0])
                ws.cell(row=i, column=2, value=aluno[1])
                ws.cell(row=i, column=3, value=aluno[2])
                ws.cell(row=i, column=4, value=aluno[3])
                ws.cell(row=i, column=5, value=aluno[4])
                                    
                foto_path = os.path.join(base_dir, aluno[5]) if aluno[5] else os.path.join(base_dir, "imagens", "00000-sem imagem.jpg")
                if os.path.exists(foto_path):
                    img = Image.open(foto_path)
                    img = img.resize((100, 100), Image.Resampling.LANCZOS)
                    temp_path = os.path.join(temp_dir, f"temp_image_{i}.jpg")
                    img.save(temp_path)
                    excel_img = XLImage(temp_path)
                    excel_img.anchor = f"F{i}"
                    ws.add_image(excel_img)
                    os.remove(temp_path)

                ws.row_dimensions[i].height = 80

            wb.save(excel_path)
            self.view.fechar_carregando()
            self.view.mostrar_mensagem("Sucesso", f"Carômetro exportado como {nome_arquivo} na pasta 'documentos'!")
            os.startfile(excel_path)
            self.view.tela_inicial()
        except Exception as e:
            self.view.fechar_carregando()
            self.view.mostrar_erro(f"Erro ao exportar Excel: {e}")

    def exportar_word(self):
        try:
            print("Debug: Iniciando exportar_word")
            turma_str = self.view.current_frame.winfo_children()[2].get()
            print(f"Debug: turma_str = {turma_str}")
            if not turma_str:
                self.view.fechar_carregando()
                self.view.mostrar_aviso("Selecione uma turma!")
                return
            turma_id = None if turma_str == "Todas" else int(turma_str.split(" - ")[0])
            print(f"Debug: turma_id = {turma_id}")
            alunos = self.model.carregar_alunos_por_turma(turma_id)
            print(f"Debug: Alunos retornados = {alunos}")
            if not alunos:
                self.view.fechar_carregando()
                self.view.mostrar_mensagem("Info", "Nenhum aluno encontrado para a turma selecionada.")
                return

            agora = datetime.now()
            filtro = "todas" if turma_str == "Todas" else turma_str.split(" - ")[1].replace(" ", "_").lower()
            nome_arquivo = f"carometro-{filtro}-{agora.strftime('%Y-%m-%d_%H-%M-%S')}.docx"
            documentos_dir = os.path.join(os.path.dirname(__file__), "documentos")
            os.makedirs(documentos_dir, exist_ok=True)
            word_path = os.path.join(documentos_dir, nome_arquivo)

            doc = Document()
            section = doc.sections[0]
            section.left_margin = Inches(0.35)
            section.right_margin = Inches(0.35)
            section.top_margin = Inches(0.79)
            section.bottom_margin = Inches(1.18)

            doc.add_heading(f"Carômetro - {turma_str}", 0)
            doc.add_paragraph()

            table = None
            row = None
            base_dir = os.path.dirname(__file__)
            for i, aluno in enumerate(alunos):
                print(f"Debug: Processando aluno = {aluno}")
                if i % 6 == 0:
                    if table:
                        doc.add_paragraph()
                    table = doc.add_table(rows=2, cols=6)
                    table.style = 'Table Grid'
                    for cell in table.rows[0].cells:
                        cell.width = Inches(1.26)
                    for cell in table.rows[1].cells:
                        cell.width = Inches(1.26)
                    row = 0
                
                foto_path = os.path.join(base_dir, aluno[5]) if aluno[5] else os.path.join(base_dir, "imagens", "00000-sem imagem.jpg")
                print(f"Debug: foto_path = {foto_path}")
                cell_img = table.rows[0].cells[i % 6]
                cell_text = table.rows[1].cells[i % 6]
                if os.path.exists(foto_path):
                    try:
                        cell_img.paragraphs[0].add_run().add_picture(foto_path, width=Inches(1.1), height=Inches(1.1))
                        cell_img.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    except Exception as e:
                        print(f"Debug: Erro ao carregar foto no Word {foto_path}: {e}")
                        cell_img.paragraphs[0].add_run("Sem Foto")
                        cell_img.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    print(f"Debug: Foto não encontrada: {foto_path}")
                    cell_img.paragraphs[0].add_run("Sem Foto")
                    cell_img.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                primeiro_nome = aluno[1].split(" ")[0]
                cell_text.paragraphs[0].add_run(primeiro_nome)
                cell_text.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell_text.paragraphs[0].runs[0].font.size = Pt(8)

                if (i + 1) % 42 == 0 and i != len(alunos) - 1:
                    doc.add_page_break()

            for section in doc.sections:
                footer = section.footer
                footer_paragraph = footer.paragraphs[0]
                footer_paragraph.text = f"Página {footer_paragraph.add_run().add_break()} " \
                                        f"Documento gerado em {agora.strftime('%d/%m/%Y, às %H:%M')}"
                footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                footer_paragraph.runs[0].font.size = Pt(8)

            doc.save(word_path)

            self.view.fechar_carregando()
            self.view.mostrar_mensagem("Sucesso", f"Carômetro exportado como {nome_arquivo} na pasta 'documentos'!")
            os.startfile(word_path)
            self.view.tela_inicial()
        except Exception as e:
            self.view.fechar_carregando()
            self.view.mostrar_erro(f"Erro ao exportar Word: {e}")

    def tela_inicial(self):
        frame = self.view.novo_frame()
        tk.Label(frame, text="Bem-vindo ao Sistema de Controle de Alunos", font=("Arial", 20, "bold"), bg="#FFFFFF", fg="#1A1A1A").pack(pady=20)
        tk.Label(frame, text="Este sistema permite gerenciar instituições, professores, cursos, turmas e alunos de forma eficiente.\nUtilize os menus acima para cadastrar, consultar e exportar dados.", font=("Arial", 10), bg="#FFFFFF", fg="#1A1A1A", justify="center").pack(pady=10)

        stats_frame = tk.Frame(frame, bg="#F8F9FA", bd=2, relief="groove")
        stats_frame.pack(pady=10, padx=10, fill=tk.X)

        inst_label = tk.Label(stats_frame, text="Instituições: Carregando...", font=("Arial", 10), bg="#F8F9FA", fg="#1A1A1A", padx=10, pady=5)
        inst_label.pack(side=tk.LEFT)
        prof_label = tk.Label(stats_frame, text="Professores: Carregando...", font=("Arial", 10), bg="#F8F9FA", fg="#1A1A1A", padx=10, pady=5)
        prof_label.pack(side=tk.LEFT)
        curso_label = tk.Label(stats_frame, text="Cursos: Carregando...", font=("Arial", 10), bg="#F8F9FA", fg="#1A1A1A", padx=10, pady=5)
        curso_label.pack(side=tk.LEFT)
        turma_label = tk.Label(stats_frame, text="Turmas: Carregando...", font=("Arial", 10), bg="#F8F9FA", fg="#1A1A1A", padx=10, pady=5)
        turma_label.pack(side=tk.LEFT)
        aluno_label = tk.Label(stats_frame, text="Alunos: Carregando...", font=("Arial", 10), bg="#F8F9FA", fg="#1A1A1A", padx=10, pady=5)
        aluno_label.pack(side=tk.LEFT)

        def carregar_contagens():
            try:
                inst_count = len(self.model.carregar_instituicoes())
                prof_count = len(self.model.carregar_professores())
                curso_count = len(self.model.carregar_cursos())
                turma_count = len(self.model.carregar_turmas())
                aluno_count = len(self.model.carregar_alunos())
                print(f"Debug: Contagens - Instituições={inst_count}, Professores={prof_count}, Cursos={curso_count}, Turmas={turma_count}, Alunos={aluno_count}")

                self.view.root.after(0, lambda: [
                    inst_label.config(text=f"Instituições: {inst_count}"),
                    prof_label.config(text=f"Professores: {prof_count}"),
                    curso_label.config(text=f"Cursos: {curso_count}"),
                    turma_label.config(text=f"Turmas: {turma_count}"),
                    aluno_label.config(text=f"Alunos: {aluno_count}")
                ])
            except Exception as e:
                print(f"Debug: Erro ao carregar contagens: {e}")
                self.view.root.after(0, lambda: [
                    inst_label.config(text="Instituições: Erro"),
                    prof_label.config(text="Professores: Erro"),
                    curso_label.config(text="Cursos: Erro"),
                    turma_label.config(text="Turmas: Erro"),
                    aluno_label.config(text="Alunos: Erro")
                ])

        carregar_contagens()  # Executar no thread principal